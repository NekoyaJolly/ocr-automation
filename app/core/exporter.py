"""各種フォーマット（TXT, Word, Excel, PDF）への出力を行うエクスポーターモジュール。"""

import logging
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Any

from app.exceptions import ExportError
from app.models.template_model import Template

logger = logging.getLogger(__name__)

# サードパーティライブラリの遅延インポート
try:
    import docx
except ImportError:
    docx = None

try:
    import openpyxl
except ImportError:
    openpyxl = None

try:
    import reportlab
except ImportError:
    reportlab = None


class Exporter(ABC):
    """ファイル出力を行うエクスポーターの抽象基底クラス。"""

    @abstractmethod
    def export(
        self, mapped_data: dict[str, Any], template: Template, output_path: Path
    ) -> None:
        """データを指定のフォーマットでファイルに書き出す。

        Args:
            mapped_data: 抽出済みのデータ（項目名 -> 値）
            template: 適用するテンプレート定義
            output_path: 書き出し先のファイルパス

        Raises:
            ExportError: 出力処理中にエラーが発生した場合
        """
        ...

    def _find_template_file(self, filename: str) -> Path | None:
        """指定されたテンプレートファイルを探索する。

        1. ユーザー定義テンプレートディレクトリ
        2. アプリケーション内蔵の templates ディレクトリ
        の順で検索する。
        """
        from app.infrastructure.paths import get_user_templates_dir

        # ユーザー設定領域の templates
        p1 = get_user_templates_dir() / filename
        if p1.exists():
            return p1

        # プロジェクトルートの templates
        p2 = Path(__file__).resolve().parent.parent.parent / "templates" / filename
        if p2.exists():
            return p2

        return None


class TxtExporter(Exporter):
    """プレーンテキスト形式でデータを出力するエクスポーター。"""

    def export(
        self, mapped_data: dict[str, Any], template: Template, output_path: Path
    ) -> None:
        try:
            lines = []
            for field in template.fields:
                val = mapped_data.get(field.output_label)
                val_str = str(val) if val is not None else ""
                lines.append(f"{field.output_label}: {val_str}")

            output_path.write_text("\n".join(lines), encoding="utf-8")
        except Exception as e:
            raise ExportError(f"テキスト出力に失敗しました: {e}") from e


class DocxExporter(Exporter):
    """Word (.docx) 形式でデータを出力するエクスポーター。"""

    def export(
        self, mapped_data: dict[str, Any], template: Template, output_path: Path
    ) -> None:
        if docx is None:
            raise ExportError(
                "python-docx がインストールされていないため、Word 出力は利用できません。"
            )

        try:
            doc = None
            if template.template_file:
                temp_path = self._find_template_file(template.template_file)
                if temp_path:
                    doc = docx.Document(temp_path)
                else:
                    logger.warning(
                        f"テンプレートファイルが見つかりません: {template.template_file}。新規作成します。"
                    )

            if doc is None:
                doc = docx.Document()
                doc.add_heading(template.name, level=1)
                for field in template.fields:
                    val = mapped_data.get(field.output_label)
                    val_str = str(val) if val is not None else ""
                    doc.add_paragraph(f"{field.output_label}: {val_str}")
            else:
                # プレースホルダーの置換処理
                for p in doc.paragraphs:
                    self._replace_placeholders(p, mapped_data)
                for table in doc.tables:
                    for row in table.rows:
                      for cell in row.cells:
                          for p in cell.paragraphs:
                              self._replace_placeholders(p, mapped_data)

            doc.save(output_path)
        except Exception as e:
            raise ExportError(f"Word 出力に失敗しました: {e}") from e

    def _replace_placeholders(self, paragraph: Any, mapped_data: dict[str, Any]) -> None:
        """段落内の {{項目名}} を実際の値に置換する。"""
        text = paragraph.text
        replaced = False
        for key, val in mapped_data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in text:
                text = text.replace(placeholder, str(val) if val is not None else "")
                replaced = True
        if replaced:
            paragraph.text = text


class XlsxExporter(Exporter):
    """Excel (.xlsx) 形式でデータを出力するエクスポーター。"""

    def export(
        self, mapped_data: dict[str, Any], template: Template, output_path: Path
    ) -> None:
        if openpyxl is None:
            raise ExportError(
                "openpyxl がインストールされていないため、Excel 出力は利用できません。"
            )

        try:
            wb = None
            if template.template_file:
                temp_path = self._find_template_file(template.template_file)
                if temp_path:
                    wb = openpyxl.load_workbook(temp_path)
                else:
                    logger.warning(
                        f"テンプレートファイルが見つかりません: {template.template_file}。新規作成します。"
                    )

            if wb is None:
                wb = openpyxl.Workbook()
                ws = wb.active
                ws.title = "OCR 結果"
                ws.cell(row=1, column=1, value="項目名")
                ws.cell(row=1, column=2, value="値")
                for idx, field in enumerate(template.fields, start=2):
                    ws.cell(row=idx, column=1, value=field.output_label)
                    val = mapped_data.get(field.output_label)
                    ws.cell(row=idx, column=2, value=val)
            else:
                ws = wb.active
                for field in template.fields:
                    val = mapped_data.get(field.output_label)
                    if field.target_position:
                        try:
                            # 'B2' などのセル番地に書き込む
                            ws[field.target_position] = val
                        except Exception as e:
                            logger.error(
                                f"Excel セル書き込み失敗 ({field.target_position}): {e}"
                            )

            wb.save(output_path)
        except Exception as e:
            raise ExportError(f"Excel 出力に失敗しました: {e}") from e


class PdfExporter(Exporter):
    """PDF 形式でデータを出力するエクスポーター。"""

    def export(
        self, mapped_data: dict[str, Any], template: Template, output_path: Path
    ) -> None:
        if reportlab is None:
            raise ExportError(
                "reportlab がインストールされていないため、PDF 出力は利用できません。"
            )

        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.cidfonts import CIDFont
            from reportlab.pdfgen import canvas

            # 日本語用フォント登録
            font_name = "Helvetica"
            try:
                pdfmetrics.registerFont(CIDFont("HeiseiKakuGo-W5"))
                font_name = "HeiseiKakuGo-W5"
            except Exception:
                logger.warning(
                    "HeiseiKakuGo-W5 日本語フォント登録に失敗しました。英語フォントで描画します。"
                )

            c = canvas.Canvas(str(output_path), pagesize=A4)
            width, height = A4

            # ヘッダータイトル
            c.setFont(font_name, 18)
            c.drawString(50, height - 50, template.name)

            # 各項目を描画
            c.setFont(font_name, 12)
            y = height - 100
            for field in template.fields:
                val = mapped_data.get(field.output_label)
                val_str = str(val) if val is not None else ""
                c.drawString(50, y, f"{field.output_label}: {val_str}")
                y -= 25
                if y < 50:
                    c.showPage()
                    c.setFont(font_name, 12)
                    y = height - 50

            c.save()
        except Exception as e:
            raise ExportError(f"PDF 出力に失敗しました: {e}") from e


class ExporterFactory:
    """出力フォーマット名に基づいて適切なエクスポーターインスタンスを生成するファクトリ。"""

    @staticmethod
    def create(format_name: str) -> Exporter:
        match format_name.lower():
            case "txt":
                return TxtExporter()
            case "docx":
                return DocxExporter()
            case "xlsx":
                return XlsxExporter()
            case "pdf":
                return PdfExporter()
            case _:
                raise ExportError(f"サポートされていない出力形式です: {format_name}")
